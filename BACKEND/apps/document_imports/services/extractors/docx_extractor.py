from pathlib import Path
from typing import Any

from docx import Document

from .base_extractor import (
    BaseExtractor,
    ExtractedPage,
    ExtractionResult,
)


class DOCXExtractor(BaseExtractor):
    """
    Extracteur de documents Microsoft Word (.docx).

    Extrait :
    - les paragraphes ;
    - les tableaux ;
    - les métadonnées du document.
    """

    SUPPORTED_EXTENSIONS = {
        ".docx",
    }

    def supports(
        self,
        file_path: str,
    ) -> bool:
        """
        Vérifie si le fichier est un document DOCX.
        """

        extension = (
            Path(file_path)
            .suffix
            .lower()
        )

        return extension in self.SUPPORTED_EXTENSIONS

    def extract(
        self,
        file_path: str,
    ) -> ExtractionResult:
        """
        Extrait le contenu du document DOCX.
        """

        if not self.supports(file_path):
            return ExtractionResult(
                success=False,
                error="Le fichier fourni n'est pas un document DOCX.",
            )

        try:
            document = Document(file_path)

            pages = []

            # Un DOCX n'a pas une notion native de "page"
            # comme un PDF. Nous considérons donc le document
            # comme une première unité d'extraction.
            text = self._extract_paragraphs(document)

            tables = self._extract_tables(document)

            page = ExtractedPage(
                page_number=1,
                text=text,
                tables=tables,
                metadata={
                    "source_type": "docx",
                    "paragraph_count": len(
                        document.paragraphs
                    ),
                    "table_count": len(
                        document.tables
                    ),
                },
            )

            pages.append(page)

            metadata = self._extract_metadata(
                document
            )

            return ExtractionResult(
                success=True,
                pages=pages,
                page_count=1,
                metadata=metadata,
            )

        except Exception as exc:

            return ExtractionResult(
                success=False,
                error=str(exc),
            )

    def _extract_paragraphs(
        self,
        document: Document,
    ) -> str:
        """
        Extrait tous les paragraphes du document.
        """

        paragraphs = []

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        return "\n".join(paragraphs)

    def _extract_tables(
        self,
        document: Document,
    ) -> list[dict[str, Any]]:
        """
        Extrait les tableaux du document.
        """

        tables = []

        for table_index, table in enumerate(
            document.tables,
            start=1,
        ):

            rows = []

            for row in table.rows:

                cells = [
                    cell.text.strip()
                    for cell in row.cells
                ]

                rows.append(cells)

            headers = []

            if rows:
                headers = rows[0]

            data_rows = (
                rows[1:]
                if len(rows) > 1
                else []
            )

            tables.append(
                {
                    "table_index": table_index,
                    "headers": headers,
                    "rows": data_rows,
                    "raw_data": {
                        "rows": rows,
                    },
                }
            )

        return tables

    def _extract_metadata(
        self,
        document: Document,
    ) -> dict[str, Any]:
        """
        Extrait les métadonnées du document Word.
        """

        properties = document.core_properties

        return {
            "title": properties.title or "",
            "author": properties.author or "",
            "subject": properties.subject or "",
            "keywords": properties.keywords or "",
            "comments": properties.comments or "",
            "category": properties.category or "",
            "created": (
                properties.created.isoformat()
                if properties.created
                else None
            ),
            "modified": (
                properties.modified.isoformat()
                if properties.modified
                else None
            ),
        }